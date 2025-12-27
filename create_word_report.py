from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# Create a new Document
doc = Document()

# Add title
title = doc.add_heading('Apple Inc. Financial Analysis Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add subtitle with period
subtitle = doc.add_paragraph('FY25 Q2 - Consolidated Financial Statements Analysis')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(14)
subtitle_format.font.italic = True

# Add date
date_para = doc.add_paragraph(f'Report Generated: {datetime.datetime.now().strftime("%B %d, %Y")}')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # Spacing

# ============ EXECUTIVE SUMMARY ============
doc.add_heading('Executive Summary', level=1)
doc.add_paragraph(
    'Apple Inc. demonstrates strong financial performance in FY25 Q2 with solid revenue growth, '
    'excellent profitability margins, and strategic diversification into high-growth services. '
    'The company maintains a robust balance sheet with substantial liquid assets despite aggressive '
    'working capital management. This report provides a comprehensive analysis of key performance '
    'indicators (KPIs) and investment recommendations.'
)

# ============ KEY METRICS OVERVIEW ============
doc.add_heading('Key Performance Indicators (KPIs)', level=1)

# Create table for KPIs
table = doc.add_table(rows=8, cols=3)
table.style = 'Light Grid Accent 1'

# Header row
header_cells = table.rows[0].cells
header_cells[0].text = 'KPI Metric'
header_cells[1].text = 'Value'
header_cells[2].text = 'Rating'

# Data rows
kpi_data = [
    ('Revenue Growth (YoY)', '4.44%', 'Healthy'),
    ('Gross Margin', '46.96%', 'Excellent'),
    ('Operating Margin', '32.97%', 'Strong'),
    ('Net Profit Margin', '27.82%', 'Excellent'),
    ('Return on Equity (Annualized)', '113.07%', 'Strong'),
    ('Debt-to-Equity Ratio', '2.06x', 'Moderate'),
    ('Current Ratio', '0.82x', 'Weak'),
]

for i, (metric, value, rating) in enumerate(kpi_data, start=1):
    row = table.rows[i].cells
    row[0].text = metric
    row[1].text = value
    row[2].text = rating

doc.add_paragraph()

# ============ FINANCIAL METRICS ============
doc.add_heading('Detailed Financial Analysis', level=1)

doc.add_heading('1. Income Statement Metrics (6-Month Period)', level=2)
doc.add_paragraph('Period: March 29, 2025 vs March 30, 2024').runs[0].font.italic = True

doc.add_paragraph().add_run('Revenue Metrics:').bold = True
revenue_points = [
    ('Total Net Sales (6M 2025)', '$219,659 Million'),
    ('Total Net Sales (6M 2024)', '$210,328 Million'),
    ('Year-over-Year Growth', '4.44%'),
    ('Products Revenue (6M 2025)', '$166,674 Million'),
    ('Services Revenue (6M 2025)', '$52,985 Million'),
    ('Product Growth (YoY)', '2.04%'),
    ('Services Growth (YoY)', '12.77%'),
]

for label, value in revenue_points:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{label}: ').bold = True
    p.add_run(value)

doc.add_paragraph().add_run('Profitability Metrics:').bold = True
profit_points = [
    ('Gross Profit (6M 2025)', '$103,142 Million'),
    ('Gross Margin (6M 2025)', '46.96%'),
    ('Gross Margin (6M 2024)', '46.18%'),
    ('Margin Improvement', '+0.78%'),
    ('Operating Income (6M 2025)', '$72,421 Million'),
    ('Operating Margin (6M 2025)', '32.97%'),
    ('Net Income (6M 2025)', '$61,110 Million'),
    ('Net Profit Margin (6M 2025)', '27.82%'),
]

for label, value in profit_points:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{label}: ').bold = True
    p.add_run(value)

doc.add_heading('2. Balance Sheet Metrics (As of March 29, 2025)', level=2)

doc.add_paragraph().add_run('Asset Metrics:').bold = True
asset_points = [
    ('Total Assets', '$331,233 Million'),
    ('Current Assets', '$118,674 Million'),
    ('Cash & Cash Equivalents', '$28,162 Million'),
    ('Marketable Securities', '$104,760 Million'),
    ('Total Liquid Assets', '$132,922 Million'),
    ('Non-Current Assets', '$212,559 Million'),
]

for label, value in asset_points:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{label}: ').bold = True
    p.add_run(value)

doc.add_paragraph().add_run('Liability & Equity Metrics:').bold = True
liability_points = [
    ('Total Liabilities', '$223,137 Million'),
    ('Current Liabilities', '$144,571 Million'),
    ('Non-Current Liabilities', '$78,566 Million'),
    ('Shareholders\' Equity', '$108,096 Million'),
    ('Debt-to-Equity Ratio', '2.06x'),
]

for label, value in liability_points:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{label}: ').bold = True
    p.add_run(value)

doc.add_paragraph().add_run('Liquidity Analysis:').bold = True
liquidity_points = [
    ('Current Ratio', '0.82x'),
    ('Quick Ratio', '0.78x'),
    ('Working Capital', '-$25,897 Million'),
]

for label, value in liquidity_points:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{label}: ').bold = True
    p.add_run(value)

doc.add_heading('3. Return & Efficiency Metrics', level=2)

doc.add_paragraph().add_run('Return Metrics (6-Month Annualized):').bold = True
return_points = [
    ('Return on Assets (ROA)', '36.90%'),
    ('Return on Equity (ROE)', '113.07%'),
    ('Asset Turnover Ratio', '0.66x'),
]

for label, value in return_points:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{label}: ').bold = True
    p.add_run(value)

doc.add_paragraph().add_run('Operating Efficiency:').bold = True
efficiency_points = [
    ('R&D Expense Ratio', '7.66%'),
    ('SG&A Expense Ratio', '6.33%'),
    ('Total Operating Expense Ratio', '13.99%'),
]

for label, value in efficiency_points:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{label}: ').bold = True
    p.add_run(value)

doc.add_paragraph()

# ============ STRENGTHS ============
doc.add_heading('Strengths', level=1)

strengths = [
    ('Strong Revenue Growth', 
     '4.44% year-over-year growth with services accelerating at 12.77%, demonstrating successful '
     'business model diversification away from hardware dependence.'),
    
    ('Excellent Profitability Margins',
     'Gross margin of 46.96%, operating margin of 32.97%, and net margin of 27.82% represent '
     'industry-leading profitability. All margins expanded year-over-year.'),
    
    ('Strong Operational Performance',
     'Operating income increased 6.04% YoY to $72.4B, and net income grew 6.20% YoY to $61.1B, '
     'demonstrating consistent earnings growth.'),
    
    ('Solid Balance Sheet',
     'Strong liquid position with $132.9B in cash and marketable securities provides flexibility '
     'for R&D investment, dividends, and strategic acquisitions.'),
    
    ('High Capital Efficiency',
     'Annualized ROE of 113.07% and ROA of 36.90% demonstrate exceptional capital efficiency and '
     'value creation for shareholders.'),
    
    ('Controlled Debt Levels',
     'Debt-to-equity ratio of 2.06x, while moderate, is sustainable given strong cash flows and '
     'excellent profitability metrics.'),
]

for title, description in strengths:
    p = doc.add_paragraph(style='List Number')
    p.add_run(f'{title}: ').bold = True
    p.add_run(description)

doc.add_paragraph()

# ============ AREAS OF CONCERN ============
doc.add_heading('Areas of Concern', level=1)

concerns = [
    ('Product Revenue Growth Slowing',
     'Core product revenue growth of only 2.04% suggests potential market saturation. Close monitoring '
     'of hardware category performance is essential to identify emerging trends.'),
    
    ('Declining Total Assets',
     'Total assets decreased 9.25% from September 2024 ($365B) to March 2025 ($331B). While likely '
     'due to intentional capital redeployment, the trend warrants monitoring.'),
    
    ('Current Ratio Below 1.0',
     'Current ratio of 0.82x indicates current liabilities exceed current assets. This is common for '
     'tech companies but requires careful working capital management.'),
    
    ('Services Profitability Metrics Unclear',
     'While services revenue is growing rapidly at 12.77%, the profitability metrics of this segment '
     'are not separately disclosed in the financial statements.'),
]

for title, description in concerns:
    p = doc.add_paragraph(style='List Number')
    p.add_run(f'{title}: ').bold = True
    p.add_run(description)

doc.add_paragraph()

# ============ INVESTMENT RECOMMENDATIONS ============
doc.add_heading('Investment Recommendations', level=1)

doc.add_heading('Overall Recommendation: BUY', level=2)
recommendation_para = doc.add_paragraph(
    'Based on the comprehensive financial analysis, Apple Inc. presents a compelling investment '
    'opportunity for long-term investors seeking exposure to a mature technology company with strong '
    'profitability and growth potential.'
)
recommendation_para.runs[0].font.bold = True
recommendation_para.runs[0].font.color.rgb = RGBColor(0, 128, 0)

doc.add_heading('1. For Long-Term Investors', level=2)
doc.add_paragraph(
    'Apple demonstrates operational excellence with improving margins, strong cash generation, and '
    'strategic diversification into high-growth services. The company\'s strong balance sheet and cash '
    'generation capability support sustainable dividends and growth investments.'
).runs[0].font.italic = True

positive_signals = [
    'Strong fundamentals with industry-leading profit margins (46.96% gross margin)',
    'Consistent revenue growth and profitability improvements year-over-year',
    'Services revenue acceleration (12.77%) suggests successful business model evolution',
    'High ROE (113.07%) indicates efficient capital deployment',
    'Ample cash reserves ($132.9B liquid assets) for R&D, dividends, and buybacks',
    'History of returning capital to shareholders through dividends and buybacks',
]

for signal in positive_signals:
    doc.add_paragraph(signal, style='List Bullet')

doc.add_heading('2. For Growth Investors', level=2)
doc.add_paragraph(
    'While total revenue growth of 4.44% is moderate, the acceleration in services (12.77%) provides '
    'a growth vector. The 2.04% product growth suggests potential market maturity requiring careful '
    'valuation assessment.'
).runs[0].font.italic = True

doc.add_heading('3. For Value Investors', level=2)
doc.add_paragraph(
    'Valuation assessment requires comparison of current P/E ratio against historical averages and '
    'industry peers. The exceptional profitability metrics and cash generation support investment at '
    'reasonable valuations.'
).runs[0].font.italic = True

doc.add_paragraph()

# ============ AREAS TO MONITOR ============
doc.add_heading('Critical Areas to Monitor', level=1)

monitoring_items = [
    ('Product Revenue Trajectory',
     'Track quarterly product revenue growth to detect potential market saturation. Look for stabilization '
     'or acceleration as key metrics.'),
    
    ('Services Segment Performance',
     'Request or identify separate services profitability metrics. Ensure growth is profitable and not '
     'cannibalizing hardware margins.'),
    
    ('Capital Allocation Decisions',
     'Monitor management\'s capital deployment strategies. Track dividend/buyback activity, R&D spending, '
     'and potential acquisitions.'),
    
    ('Macroeconomic Headwinds',
     'Watch for impacts of consumer spending slowdowns, currency fluctuations (significant international '
     'presence), and inflation on margins.'),
    
    ('Regulatory Developments',
     'Monitor antitrust investigations and App Store policy challenges, particularly in key markets like EU.'),
    
    ('Working Capital Management',
     'Track the current ratio trend. While current ratio of 0.82x is manageable, monitor for deterioration.'),
]

for title, description in monitoring_items:
    p = doc.add_paragraph(style='List Number')
    p.add_run(f'{title}: ').bold = True
    p.add_run(description)

doc.add_paragraph()

# ============ ACTION ITEMS ============
doc.add_heading('Recommended Action Items', level=1)

action_items = [
    'Compare Apple\'s P/E ratio (price-to-earnings) with S&P 500 average (~20-22x) and industry peers',
    'Review Apple\'s latest quarterly guidance and management commentary on product/services outlook',
    'Analyze historical dividend yield and capital return history to assess shareholder-friendly policies',
    'Compare gross and operating margins with key competitors (Microsoft, Google, Samsung, etc.)',
    'Evaluate iPhone, Mac, Services, and Wearables segment performance separately',
    'Monitor Apple\'s R&D spending trends and innovation pipeline',
    'Assess macroeconomic indicators (consumer spending, currency trends, interest rates) for impact',
    'Review analyst consensus price targets and earnings estimates for next 2-4 quarters',
]

for i, item in enumerate(action_items, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(item)

doc.add_paragraph()

# ============ INVESTMENT THESIS ============
doc.add_heading('Investment Thesis Summary', level=1)

thesis_points = [
    ('Operational Excellence',
     'Apple continues to demonstrate superior operational execution with expanding margins despite slower growth.'),
    
    ('Business Model Diversification',
     'Services growth at 12.77% provides earnings diversification and reduces hardware dependency.'),
    
    ('Capital Efficiency',
     'ROE of 113.07% and ROA of 36.90% demonstrate exceptional capital efficiency.'),
    
    ('Financial Strength',
     'Robust balance sheet with $132.9B in liquid assets supports sustainable operations and shareholder returns.'),
    
    ('Strategic Positioning',
     'Strong brand value, customer loyalty, and ecosystem lock-in provide competitive moat.'),
]

for title, description in thesis_points:
    p = doc.add_paragraph(style='List Number')
    p.add_run(f'{title}: ').bold = True
    p.add_run(description)

doc.add_paragraph()

# ============ RISK ASSESSMENT ============
doc.add_heading('Risk Assessment', level=1)

overall_risk = doc.add_paragraph()
overall_risk.add_run('Overall Risk Level: ').bold = True
overall_risk.add_run('MODERATE').bold = True
overall_risk.runs[-1].font.color.rgb = RGBColor(255, 165, 0)

risk_factors = [
    ('Product Market Maturity',
     'Slower product growth (2.04%) suggests potential market saturation, particularly in developed markets.'),
    
    ('Macroeconomic Sensitivity',
     'Consumer discretionary spending vulnerability to economic downturns affects hardware sales.'),
    
    ('Currency Risk',
     'Significant international revenue exposure creates currency fluctuation risk.'),
    
    ('Regulatory Risk',
     'Antitrust investigations and App Store policy challenges could impact profitability.'),
    
    ('Competitive Pressure',
     'Intense competition in smartphones, wearables, and services from Samsung, Google, Amazon, etc.'),
]

for title, description in risk_factors:
    p = doc.add_paragraph(style='List Number')
    p.add_run(f'{title}: ').bold = True
    p.add_run(description)

doc.add_paragraph()

# ============ CONCLUSION ============
doc.add_heading('Conclusion', level=1)

conclusion = (
    'Apple Inc. demonstrates strong financial fundamentals with excellent profitability metrics, solid revenue '
    'growth, and strategic business model diversification through services. The company\'s $132.9B in liquid assets, '
    'exceptional ROE of 113.07%, and consistent margin expansion provide confidence in financial strength. '
    '\n\n'
    'While product revenue growth is moderating at 2.04%, the acceleration in services at 12.77% suggests successful '
    'diversification. The moderate risk profile, combined with strong operational performance, makes Apple an attractive '
    'investment for long-term investors at reasonable valuations.\n\n'
    'Key success factors to monitor include: (1) product revenue stabilization, (2) services profitability verification, '
    '(3) capital allocation strategy, and (4) macroeconomic headwinds.\n\n'
    'Recommendation: BUY for long-term investors seeking quality, profitability, and growth with moderate risk exposure.'
)

doc.add_paragraph(conclusion)

doc.add_paragraph()

# ============ FOOTER ============
footer = doc.add_paragraph('_' * 90)
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

footer_text = doc.add_paragraph(
    f'Report Generated: {datetime.datetime.now().strftime("%B %d, %Y at %I:%M %p")}\n'
    'Company: Apple Inc. | Period: FY25 Q2 (6 months ended March 29, 2025)\n'
    'Source: FY25 Q2 Consolidated Financial Statements'
)
footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_text.runs[0].font.size = Pt(9)
footer_text.runs[0].font.color.rgb = RGBColor(128, 128, 128)

# Save the document
doc.save('Apple_Financial_Analysis_Report.docx')
print("✅ Word document created successfully: Apple_Financial_Analysis_Report.docx")
